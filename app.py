# from flask import Flask, request, render_template, jsonify
# import openai
# import pandas as pd
# import PyPDF2
# from docx import Document
# import os

# # Настройка API ключа OpenAI
# openai.api_key = "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"

# # Создаем экземпляр Flask
# app = Flask(__name__)

# # Функция для обработки Excel файлов
# def process_excel(file_path):
#     try:
#         data = pd.read_excel(file_path, engine='openpyxl')
#         text_data = data.to_string()
#         return text_data
#     except Exception as e:
#         return f"Error processing Excel file: {e}"

# # Функция для обработки CSV файлов
# def process_csv(file_path):
#     try:
#         data = pd.read_csv(file_path)
#         text_data = data.to_string()
#         return text_data
#     except Exception as e:
#         return f"Error processing CSV file: {e}"

# # Функция для обработки PDF файлов
# def process_pdf(file_path):
#     try:
#         with open(file_path, 'rb') as file:
#             reader = PyPDF2.PdfReader(file)
#             text_data = ""
#             for page in range(len(reader.pages)):
#                 text_data += reader.pages[page].extract_text()
#             return text_data
#     except Exception as e:
#         return f"Error processing PDF file: {e}"

# # Функция для обработки Word файлов (.docx)
# def process_word(file_path):
#     try:
#         doc = Document(file_path)
#         text_data = ""
#         for para in doc.paragraphs:
#             text_data += para.text + "\n"
#         return text_data
#     except Exception as e:
#         return f"Error processing Word file: {e}"

# # Функция для обработки текстовых файлов
# def process_text_file(file_path):
#     try:
#         with open(file_path, 'r', encoding='utf-8') as file:
#             text_data = file.read()
#         return text_data
#     except Exception as e:
#         return f"Error processing text file: {e}"

# # Функция для выполнения запроса с использованием данных файла
# def analyze_file_with_data(text_data):
#     try:
#         response = openai.ChatCompletion.create(
#             model="gpt-3.5-turbo",
#             messages=[
#                 {"role": "system", "content": "Вы — полезный ассистент, который говорит только на русском языке."},
#                 {"role": "user", "content": f"Пожалуйста, проанализируйте следующие данные на русском языке: \n{text_data}"}
#             ]
#         )
#         return response['choices'][0]['message']['content']
#     except Exception as e:
#         return f"Ошибка при анализе данных с OpenAI: {e}"

# # Главная страница, форма для загрузки файла
# @app.route('/')
# def index():
#     return render_template('index.html')

# # Обработка загрузки файла
# @app.route('/upload', methods=['POST'])
# def upload_file():
#     if 'file' not in request.files:
#         return jsonify({"error": "No file part"})
    
#     file = request.files['file']
    
#     if file.filename == '':
#         return jsonify({"error": "No selected file"})
    
#     file_path = os.path.join('uploads', file.filename)
#     file.save(file_path)

#     # Обработка файла в зависимости от его типа
#     text_data = None
#     if file.filename.endswith('.xlsx') or file.filename.endswith('.xls'):
#         text_data = process_excel(file_path)
#     elif file.filename.endswith('.csv'):
#         text_data = process_csv(file_path)
#     elif file.filename.endswith('.pdf'):
#         text_data = process_pdf(file_path)
#     elif file.filename.endswith('.docx'):
#         text_data = process_word(file_path)
#     elif file.filename.endswith('.txt'):
#         text_data = process_text_file(file_path)
#     else:
#         return jsonify({"error": "Unsupported file type"})

#     if text_data:
#         # Анализируем данные с помощью OpenAI
#         analysis = analyze_file_with_data(text_data)
#         return render_template('result.html', analysis=analysis, text_data=text_data)

#     return jsonify({"error": "Failed to process file"})

# # Запуск приложения
# if __name__ == '__main__':
#     # Убедитесь, что папка для загрузок существует
#     if not os.path.exists('uploads'):
#         os.makedirs('uploads')
#     app.run(debug=True)
from flask import Flask, request, render_template, jsonify
import openai
import pandas as pd
import PyPDF2
from docx import Document
import os
import easyocr  # Для OCR (распознавание текста на изображениях)
from PIL import Image
# Настройка API ключа OpenAI
openai.api_key = "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"

# Создаем экземпляр Flask
app = Flask(__name__)

# Функция для обработки Excel файлов
def process_excel(file_path):
    try:
        data = pd.read_excel(file_path, engine='openpyxl')
        text_data = data.to_string()
        return text_data
    except Exception as e:
        return f"Error processing Excel file: {e}"

# Функция для обработки CSV файлов
def process_csv(file_path):
    try:
        data = pd.read_csv(file_path)
        text_data = data.to_string()
        return text_data
    except Exception as e:
        return f"Error processing CSV file: {e}"

# Функция для обработки PDF файлов
def process_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text_data = ""
            for page in range(len(reader.pages)):
                text_data += reader.pages[page].extract_text()
            return text_data
    except Exception as e:
        return f"Error processing PDF file: {e}"

# Функция для обработки Word файлов (.docx)
def process_word(file_path):
    try:
        doc = Document(file_path)
        text_data = ""
        for para in doc.paragraphs:
            text_data += para.text + "\n"
        return text_data
    except Exception as e:
        return f"Error processing Word file: {e}"

# Функция для обработки текстовых файлов
def process_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text_data = file.read()
        return text_data
    except Exception as e:
        return f"Error processing text file: {e}"

# Функция для обработки изображений с использованием OCR
def process_image(file_path):
    try:
       
        reader = easyocr.Reader(['en', 'ru'])  
        result = reader.readtext(file_path) 

        if not result:
            return "Текст на изображении не был распознан."
        
        text_data = " ".join([text[1] for text in result])
        return text_data
    except Exception as e:
        return f"Ошибка при обработке изображения: {e}"
# Функция для выполнения запроса с использованием данных файла
def analyze_file_with_data(text_data):
    try:
        if len(text_data.strip()) < 10: 
            return "Пожалуйста, предоставьте больше данных для анализа."

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": "Вы — полезный ассистент, который говорит только на русском языке."},
                      {"role": "user", "content": f"Пожалуйста, проанализируйте следующие данные на русском языке: \n{text_data}"}]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        return f"Ошибка при анализе данных с OpenAI: {e}"


@app.route('/')
def index():
    return render_template('index.html')

# Обработка загрузки файла
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"})
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No selected file"})
    
    file_path = os.path.join('uploads', file.filename)
    file.save(file_path)

    # Обработка файла в зависимости от его типа
    text_data = None
    if file.filename.endswith('.xlsx') or file.filename.endswith('.xls'):
        text_data = process_excel(file_path)
    elif file.filename.endswith('.csv'):
        text_data = process_csv(file_path)
    elif file.filename.endswith('.pdf'):
        text_data = process_pdf(file_path)
    elif file.filename.endswith('.docx'):
        text_data = process_word(file_path)
    elif file.filename.endswith('.txt'):
        text_data = process_text_file(file_path)
    elif file.filename.endswith(('.jpg', '.png', '.jpeg')): 
        text_data = process_image(file_path)
    else:
        return jsonify({"error": "Unsupported file type"})

    if text_data:
      
        analysis = analyze_file_with_data(text_data)
        return render_template('result.html', analysis=analysis, text_data=text_data)

    return jsonify({"error": "Failed to process file"})

# Запуск приложения
if __name__ == '__main__':
   
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
    app.run(debug=True)

